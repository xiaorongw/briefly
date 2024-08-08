import streamlit as st
import os
from transformers import BartTokenizer, BartForConditionalGeneration
import docx
import pandas as pd

# Cache the model and tokenizer to load them only once
@st.cache_resource
def load_model_and_tokenizer():
    tokenizer = BartTokenizer.from_pretrained('C:/Users/User/Desktop/Work/file-summarise/nlp/facebook/bart-large-cnn')
    model = BartForConditionalGeneration.from_pretrained('C:/Users/User/Desktop/Work/file-summarise/nlp/facebook/bart-large-cnn')
    return tokenizer, model

tokenizer, model = load_model_and_tokenizer()

# Function to split text into chunks with overlap
def chunk_text(text, max_length, overlap=100):
    tokens = tokenizer.encode(text, return_tensors='pt')[0]
    chunks = []
    for i in range(0, len(tokens), max_length - overlap):
        chunk = tokens[i:i + max_length]
        if len(chunk) > 0:
            chunks.append(chunk)
    return chunks

# Function to summarize text
def summarize(text):
    max_length = 1024  # Maximum length of tokens
    chunks = chunk_text(text, max_length, overlap=100)
    
    summaries = []
    for chunk in chunks:
        chunk = chunk.unsqueeze(0).to(model.device)  # Move to device if using GPU
        summary_ids = model.generate(
            chunk,
            max_length=150,
            min_length=50,
            length_penalty=1.5,
            num_beams=5,
            early_stopping=True
        )
        summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        summaries.append(summary)
    
    combined_summary = ' '.join(summaries)
    
    # Optional: Log the combined summary for debugging
    print("Combined Summary:", combined_summary)
    
    return combined_summary

# Function to extract text from .docx files
def read_docx(file_path):
    doc = docx.Document(file_path)
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip() != ""]
    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip() != ""]
            if cells:
                tables.append(" | ".join(cells))
    
    text = "\n".join(paragraphs + tables)
    
    # Print out the first 1000 characters for debugging
    # print("Extracted text snippet:", text[:1000])
    
    return text


# Function to extract text from .xlsx files
def read_xlsx(file_path):
    df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
    text = []
    for sheet_name, sheet_df in df.items():
        sheet_text = sheet_df.to_string(index=False, header=False)
        text.append(f"Sheet: {sheet_name}\n{sheet_text}")
    
    # Print out the first 1000 characters for debugging
    # print("Extracted text snippet:", "\n".join(text)[:1000])
    
    return "\n".join(text)


# Streamlit UI
st.title("Local File Navigator with NLP Summary")

# Set the default directory to the user's home directory
default_directory = os.path.expanduser("~")

if 'current_dir' not in st.session_state:
    st.session_state.current_dir = default_directory

# Display current directory
st.write(f"Current directory: {st.session_state.current_dir}")

# List directories and files in the current directory
entries = os.listdir(st.session_state.current_dir)
# Filter out hidden files and folders
entries = [e for e in entries if not e.startswith('.')]

directories = [e for e in entries if os.path.isdir(os.path.join(st.session_state.current_dir, e))]
files = [e for e in entries if e.endswith(('.docx', '.xlsx'))]

# Option to navigate back
if st.session_state.current_dir != default_directory:
    if st.button("Go Up"):
        st.session_state.current_dir = os.path.dirname(st.session_state.current_dir)

# Display directories as clickable links
st.subheader("Folders:")
for directory in directories:
    if st.button(f"📁 {directory}", key=directory):
        st.session_state.current_dir = os.path.join(st.session_state.current_dir, directory)

# Display files as clickable links
st.subheader("Files:")
for file in files:
    if st.button(f"📄 {file}", key=file):
        file_path = os.path.join(st.session_state.current_dir, file)
        st.write(f"Selected file: {file}")

        # Generate summary based on file type
        if file.endswith('.docx'):
            text = read_docx(file_path)
        elif file.endswith('.xlsx'):
            text = read_xlsx(file_path)

        if text:
            # Show status message and spinner while generating summary
            with st.spinner("Generating summary..."):
                summary = summarize(text)

            st.subheader("Summary")
            st.write(summary)
